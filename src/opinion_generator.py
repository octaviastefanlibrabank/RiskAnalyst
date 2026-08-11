"""
Generates the final structured risk opinion.

Grades (risk_level) for every risk type come 100% from KoEngineResult (Python,
deterministic) - GPT is only asked to formulate the "mentions" free text and
recommendations, and is explicitly told the grades are fixed and must not be
changed. See OpinionDraft (LLM-facing, no grade field) vs RiskMention (final,
grade injected by Python).
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Optional

from pydantic import BaseModel, Field

from .azure_llm import AzureLLM
from .models import (
    MISSING_IBS_FLOW,
    MISSING_INFO,
    CompanyRiskData,
    KoEngineResult,
    RiskMention,
    RiskOpinion,
    RuleStatus,
)
from .utils import StepLogger

LEVEL_LABELS_RO = {"SCAZUT": "Risc scazut", "MEDIU": "Risc mediu", "RIDICAT": "Risc ridicat"}


class RiskMentionDraft(BaseModel):
    type: str
    mentions: str


class OpinionDraft(BaseModel):
    overall_context_note: Optional[str] = None
    recommendations: list[str] = Field(default_factory=list)
    risks: list[RiskMentionDraft] = Field(default_factory=list)


SYSTEM_PROMPT = """\
Esti un asistent care formuleaza in limba romana mentiunile pentru o opinie de \
risc bancara de tip Corporate, pe baza datelor deja extrase si a gradelor de \
risc DEJA CALCULATE de un motor determinist in Python.

Reguli OBLIGATORII:
1. NU modifica, NU recalcula si NU sugera alte grade de risc decat cele primite. \
   Gradele (scazut/mediu/ridicat) sunt finale si nu apar in raspunsul tau - \
   formulezi doar textul ("mentions") pentru fiecare tip de risc din lista primita.
2. Foloseste STRICT informatiile din datele structurate primite. Nu inventa cifre, \
   nume sau evenimente care nu apar in date.
3. Pentru orice tip de risc marcat ca NOT_IMPLEMENTED sau DATA_MISSING, mentioneaza \
   explicit ce informatie lipseste, folosind formularile primite (nu inventa o valoare).
4. Recomandarile trebuie sa fie scurte, concrete si bazate pe riscurile identificate \
   (ex: solicitare aviz AML, garantii suplimentare, clarificare rating etc.) - nu \
   sunt o decizie finala, ci sugestii pentru ofiterul de risc.
5. Nu pretinde ca iei decizia finala - opinia va fi validata de un ofiter de risc uman.
"""


def _build_llm_context(company_name: str, data: CompanyRiskData, ko: KoEngineResult) -> str:
    payload = {
        "company_data": json.loads(data.model_dump_json()),
        "ko_categories": [
            {
                "category": c.category,
                "status": c.status.value,
                "risk_level": c.risk_level,
                "completeness": c.completeness,
                "rules": [
                    {
                        "label": r.label,
                        "status": r.status.value,
                        "risk_level": r.risk_level,
                        "explanation": r.explanation,
                    }
                    for r in c.rules
                ],
            }
            for c in ko.categories
        ],
        "overall_risk_level": ko.overall_risk_level,
        "overall_completeness": ko.overall_completeness,
    }
    return f"Compania: {company_name}\n\nDate structurate + rezultate KO (JSON):\n{json.dumps(payload, ensure_ascii=False, indent=2)}"


def _fallback_mentions(category) -> str:
    """Deterministic, non-LLM mention text - used with --no-llm or if a category has no LLM text."""
    ok = [r for r in category.rules if r.status == RuleStatus.OK]
    missing = [r for r in category.rules if r.status != RuleStatus.OK]
    parts = []
    if ok:
        parts.append("; ".join(f"{r.label}: {r.risk_level}" for r in ok) + ".")
    if missing:
        parts.append("Date lipsa/reguli neimplementate: " + "; ".join(r.label for r in missing) + ".")
    return " ".join(parts) if parts else MISSING_INFO


def generate_opinion(
    company_name: str,
    data: CompanyRiskData,
    ko: KoEngineResult,
    llm: AzureLLM | None,
    logger: StepLogger | None = None,
) -> RiskOpinion:
    draft: OpinionDraft | None = None
    if llm is not None:
        context = _build_llm_context(company_name, data, ko)
        if logger:
            logger.info("Requesting mention formulation from GPT-5-mini (grades are fixed by KO engine).")
        draft = llm.extract_structured(SYSTEM_PROMPT, context, OpinionDraft)

    draft_by_type = {}
    if draft:
        for rm in draft.risks:
            draft_by_type[rm.type.strip().lower()] = rm.mentions

    risks: list[RiskMention] = []
    for cat in ko.categories:
        grade = LEVEL_LABELS_RO.get(cat.risk_level, "Risc indeterminat (date insuficiente)")
        mentions = draft_by_type.get(cat.category.strip().lower())
        if not mentions:
            mentions = _fallback_mentions(cat)
        risks.append(
            RiskMention(
                type=cat.category,
                grade=grade,
                mentions=mentions,
                missing_inputs=cat.missing_inputs,
            )
        )

    overall_risk = LEVEL_LABELS_RO.get(ko.overall_risk_level, "Risc indeterminat (date insuficiente pentru o opinie KO completa)")

    missing_fields = list(data.missing_information)
    for cat in ko.categories:
        for r in cat.rules:
            if r.status != RuleStatus.OK:
                missing_fields.append(f"{cat.category} / {r.label}: {r.explanation}")
    missing_fields.extend(ko.notes)

    recommendations = draft.recommendations if draft and draft.recommendations else [
        "Recomandari generate automat indisponibile (--no-llm) - a se completa manual de ofiterul de risc."
    ]

    return RiskOpinion(
        client=data.company_name or company_name,
        cui=data.cui or MISSING_INFO,
        branch=data.branch or MISSING_IBS_FLOW,
        overall_risk=overall_risk,
        recommendations=recommendations,
        risks=risks,
        missing_fields=missing_fields,
    )


# --------------------------------------------------------------------------- #
# Rendering
# --------------------------------------------------------------------------- #
def render_markdown(company_name: str, data: CompanyRiskData, ko: KoEngineResult, opinion: RiskOpinion) -> str:
    lines = [f"# Opinie de risc - {opinion.client}", ""]
    lines.append(f"> {opinion.disclaimer}")
    lines.append("")
    lines.append("## Date generale")
    lines.append("")
    lines.append("| Camp | Valoare |")
    lines.append("|---|---|")
    general = [
        ("Client", opinion.client),
        ("CUI", opinion.cui),
        ("Branch", opinion.branch),
        ("Categorie client", opinion.categorie_client),
        ("Id Dosar", data.id_dosar or MISSING_IBS_FLOW),
        ("Credit ID", MISSING_IBS_FLOW),
        ("CAEN", data.caen_code or MISSING_INFO),
        ("Domeniu de activitate", data.activity_description or MISSING_INFO),
        ("Rating", data.rating or MISSING_INFO),
        ("Solicitare curenta", data.current_request or MISSING_INFO),
        ("Suma solicitata", f"{data.requested_amount:,.0f} {data.currency}" if data.requested_amount and data.currency else MISSING_INFO),
        ("Durata (luni)", str(data.duration_months) if data.duration_months else MISSING_INFO),
        ("Opinie Risc (generala)", opinion.overall_risk),
    ]
    for k, v in general:
        lines.append(f"| {k} | {v} |")
    lines.append("")

    lines.append("## Recomandari")
    lines.append("")
    for rec in opinion.recommendations:
        lines.append(f"- {rec}")
    lines.append("")

    lines.append("## Riscuri")
    lines.append("")
    lines.append("| Tip risc | Grad | Mentiuni |")
    lines.append("|---|---|---|")
    for r in opinion.risks:
        mentions = r.mentions.replace("\n", " ").replace("|", "/")
        lines.append(f"| {r.type} | {r.grade} | {mentions} |")
    lines.append("")

    lines.append("## Detaliu reguli KO (transparenta calcul)")
    lines.append("")
    lines.append("| Categorie | Regula | Status | Grad | Explicatie |")
    lines.append("|---|---|---|---|---|")
    for cat in ko.categories:
        for r in cat.rules:
            expl = r.explanation.replace("\n", " ").replace("|", "/")
            lines.append(f"| {cat.category} | {r.label} | {r.status.value} | {r.risk_level or '-'} | {expl} |")
    lines.append("")

    lines.append("## Campuri lipsa / neimplementate")
    lines.append("")
    for m in opinion.missing_fields:
        lines.append(f"- {m}")
    lines.append("")

    if data.evidence:
        lines.append("## Surse (evidence)")
        lines.append("")
        lines.append("| Camp | Valoare | Document sursa | Fragment |")
        lines.append("|---|---|---|---|")
        for e in data.evidence:
            excerpt = (e.excerpt or "").replace("\n", " ").replace("|", "/")
            lines.append(f"| {e.field} | {e.value} | {e.source_document} | {excerpt} |")
        lines.append("")

    return "\n".join(lines)


def render_html(markdown_text: str, title: str) -> str:
    # Minimal, dependency-free markdown->HTML good enough for a demo view (tables + headers + lists).
    import html as _html

    def esc(s: str) -> str:
        return _html.escape(s)

    lines = markdown_text.split("\n")
    out = [f"<!doctype html><html lang='ro'><head><meta charset='utf-8'><title>{esc(title)}</title>",
           "<style>body{font-family:Arial,sans-serif;max-width:1000px;margin:2rem auto;padding:0 1rem;line-height:1.5}",
           "table{border-collapse:collapse;width:100%;margin:1rem 0}th,td{border:1px solid #ccc;padding:6px 10px;text-align:left;vertical-align:top}",
           "th{background:#f0f0f0}blockquote{color:#555;border-left:3px solid #ccc;padding-left:1rem}</style></head><body>"]

    in_table = False
    in_list = False
    for line in lines:
        if line.startswith("|") and not line.replace("|", "").replace("-", "").strip():
            continue  # markdown table separator row
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if not in_table:
                out.append("<table>")
                out.append("<tr>" + "".join(f"<th>{esc(c)}</th>" for c in cells) + "</tr>")
                in_table = True
            else:
                out.append("<tr>" + "".join(f"<td>{esc(c)}</td>" for c in cells) + "</tr>")
            continue
        elif in_table:
            out.append("</table>")
            in_table = False

        if line.startswith("- "):
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{esc(line[2:])}</li>")
            continue
        elif in_list:
            out.append("</ul>")
            in_list = False

        if line.startswith("# "):
            out.append(f"<h1>{esc(line[2:])}</h1>")
        elif line.startswith("## "):
            out.append(f"<h2>{esc(line[3:])}</h2>")
        elif line.startswith("> "):
            out.append(f"<blockquote>{esc(line[2:])}</blockquote>")
        elif line.strip():
            out.append(f"<p>{esc(line)}</p>")

    if in_table:
        out.append("</table>")
    if in_list:
        out.append("</ul>")
    out.append("</body></html>")
    return "\n".join(out)


def render_docx(company_name: str, data: CompanyRiskData, ko: KoEngineResult, opinion: RiskOpinion):
    """Builds a python-docx Document mirroring render_markdown's structure."""
    from docx import Document
    from docx.shared import Pt

    def _xml_ok(ch):
        cp = ord(ch)
        return cp in (0x09, 0x0A, 0x0D) or 0x20 <= cp <= 0xD7FF or 0xE000 <= cp <= 0xFFFD or 0x10000 <= cp <= 0x10FFFF

    def _clean(s):
        return "".join(ch for ch in str(s) if _xml_ok(ch))

    doc = Document()
    doc.add_heading(_clean(f"Opinie de risc - {opinion.client}"), level=0)

    p = doc.add_paragraph()
    run = p.add_run(_clean(opinion.disclaimer))
    run.italic = True
    run.font.size = Pt(9)

    def add_table(headers: list[str], rows: list[list[str]]):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = _clean(h)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = _clean(v)
        return t

    doc.add_heading("Date generale", level=1)
    general = [
        ("Client", opinion.client),
        ("CUI", opinion.cui),
        ("Branch", opinion.branch),
        ("Categorie client", opinion.categorie_client),
        ("Id Dosar", data.id_dosar or MISSING_IBS_FLOW),
        ("Credit ID", MISSING_IBS_FLOW),
        ("CAEN", data.caen_code or MISSING_INFO),
        ("Domeniu de activitate", data.activity_description or MISSING_INFO),
        ("Rating", data.rating or MISSING_INFO),
        ("Solicitare curenta", data.current_request or MISSING_INFO),
        ("Suma solicitata", f"{data.requested_amount:,.0f} {data.currency}" if data.requested_amount and data.currency else MISSING_INFO),
        ("Durata (luni)", str(data.duration_months) if data.duration_months else MISSING_INFO),
        ("Opinie Risc (generala)", opinion.overall_risk),
    ]
    add_table(["Camp", "Valoare"], [[k, v] for k, v in general])

    doc.add_heading("Recomandari", level=1)
    for rec in opinion.recommendations:
        doc.add_paragraph(_clean(rec), style="List Bullet")

    doc.add_heading("Riscuri", level=1)
    add_table(
        ["Tip risc", "Grad", "Mentiuni"],
        [[r.type, r.grade, r.mentions] for r in opinion.risks],
    )

    doc.add_heading("Detaliu reguli KO (transparenta calcul)", level=1)
    ko_rows = []
    for cat in ko.categories:
        for r in cat.rules:
            ko_rows.append([cat.category, r.label, r.status.value, r.risk_level or "-", r.explanation])
    add_table(["Categorie", "Regula", "Status", "Grad", "Explicatie"], ko_rows)

    doc.add_heading("Campuri lipsa / neimplementate", level=1)
    for m in opinion.missing_fields:
        doc.add_paragraph(_clean(m), style="List Bullet")

    if data.evidence:
        doc.add_heading("Surse (evidence)", level=1)
        add_table(
            ["Camp", "Valoare", "Document sursa", "Fragment"],
            [[e.field, e.value, e.source_document, e.excerpt or ""] for e in data.evidence],
        )

    return doc


def save_outputs(company_name: str, data: CompanyRiskData, ko: KoEngineResult, opinion: RiskOpinion, output_root: Path) -> Path:
    out_dir = output_root / company_name
    out_dir.mkdir(parents=True, exist_ok=True)

    (out_dir / "opinion.json").write_text(
        json.dumps(
            {
                "opinion": json.loads(opinion.model_dump_json()),
                "ko_result": json.loads(ko.model_dump_json()),
                "extracted_data": json.loads(data.model_dump_json()),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    md = render_markdown(company_name, data, ko, opinion)
    (out_dir / "opinion.md").write_text(md, encoding="utf-8")
    (out_dir / "opinion.html").write_text(render_html(md, f"Opinie risc - {opinion.client}"), encoding="utf-8")

    try:
        docx_doc = render_docx(company_name, data, ko, opinion)
        docx_doc.save(str(out_dir / "opinion.docx"))
    except PermissionError:
        # Likely open in Word/another app - don't fail the whole run over the docx.
        print(f"      ! Nu am putut suprascrie opinion.docx (probabil deschis in alta aplicatie) - json/md/html salvate oricum.")

    return out_dir
